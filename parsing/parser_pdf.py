# Импортируем модуль os для работы с файлами и папками
import os

from docx import Document
from loguru import logger
from config import ConfigSettings


# Определяем класс FileChecker
class FileChecker:
    """
    Класс для проверки типов файлов в папке и запуска разных функций в зависимости от типа.
    """

    # Определяем конструктор класса
    def __init__(self, xml_zip_dr: (str) = None, extract_dir_xml: (str) = None,
                 pdf_zip_dir: (str) = None, extract_dir_pdf: (str) = None) -> None:
        '''
        Конструктор класса Extract. Значения для аргументов заданы как необязательные,
        для применения аргументов в дочерних классах по одному
            Аргументы
                :param xml_zip_dr (str): путь к папке с архивами zip, содержащими xml-документы
                :param extract_dir_xml (str): путь к папке, куда извлекать xml-документы
                :param pdf_zip_dir (str): путь к папке с архивами zip, содержащими pdf-документы
                :param extract_dir_pdf (str): путь к папке, куда извлекать pdf-документы
        '''
        # присваиваем аргументы конструктора атрибутам класса
        self.xml_zip_dr = xml_zip_dr
        self.extract_dir_xml = extract_dir_xml
        self.pdf_zip_dir = pdf_zip_dir
        self.extract_dir_pdf = extract_dir_pdf



    # Определяем метод check_file_type
    def check_file_type(self, file_path):
        """
        Метод для проверки типа файла по его расширению и запуска соответствующей функции для поиска ключевых слов по документу.
        :param file_path: путь к файлу, который нужно проверить
        :return: результат работы соответствующей функции или None, если тип файла не поддерживается
        """
        # Получаем расширение файла из его пути
        file_extension = os.path.splitext(file_path)[1]
        # Проверяем, какое расширение имеет файл
        if file_extension == ".docx":
            # Если файл является документом Word, то запускаем функцию для поиска слов заранее определенных слов для поиска в Word-файле
            # Пока такой функции нет, пусть она будет с pass
            pass
        elif file_extension == ".xlsx":
            # Если файл является документом Excel, то запускаем функцию для работы с Excel-файлами
            # Пока такой функции нет, пусть она будет с pass
            pass
        elif file_extension == ".pdf":
            # Если файл является документом PDF, то запускаем функцию для работы с PDF-файлами
            # Эта функция уже определена выше, мы можем использовать ее
            return self.extract_search_phrase_pdf(file_path)
        else:
            # Если файл имеет другое расширение, то возвращаем None, так как тип файла не поддерживается
            return None

    # Определяем метод extract_search_phrase_pdf
    def extract_search_phrase_pdf(self, file_path):
        """
        Метод для парсинга и извлечения данных из файла PDF, сверяет на словосочетание в тексте
        документа и возвращает строку с найденным текстом и название документа.

        :param file_path: принимает путь до файла PDF
        :return: возвращает найденную строку в документе целиком, если она содержит в себе указанное слово,
         и наименование файла в котором найдено слово.
        """
        # Здесь вы можете использовать код, который я показал вам ранее, для работы с PDF-файлами
        # Я не буду повторять его здесь, чтобы сэкономить место
        pass

    def extract_search_phrase_word(self, file_path, search_terms):
        """
        Метод для парсинга и извлечения данных из файла Word, сверяет на словосочетание в тексте
        документа и возвращает строку с найденным текстом и название документа, а также данные слева и справа от найденной фразы в той же строке таблицы.

        :param file_path: принимает путь до файла Word
        :param search_terms: принимает список словосочетаний для поиска
        :return: возвращает найденную строку в документе целиком, если она содержит в себе указанное слово,
         и наименование файла в котором найдено слово, а также данные слева и справа от найденной фразы.
        """
        try:
            doc = Document(file_path)
            file_name = os.path.basename(file_path)
            found_terms = {term.lower(): [] for term in search_terms}

            # Поиск в параграфах
            for paragraph in doc.paragraphs:
                for term in found_terms:
                    if term in paragraph.text.lower():
                        found_terms[term].append(paragraph.text)

            # Поиск в таблицах
            for table in doc.tables:
                for row in table.rows:
                    cells_text = [cell.text for cell in row.cells]
                    for cell_index, cell_text in enumerate(cells_text):
                        for term in found_terms:
                            if term in cell_text.lower():
                                left_data = ' '.join(cells_text[:cell_index])  # Данные слева от фразы
                                right_data = ' '.join(cells_text[cell_index + 1:])  # Данные справа от фразы
                                found_terms[term].append(f'{left_data} {cell_text} {right_data}')
                                logger.info(f'Найдено в файле {file_name}: {left_data} "{cell_text}" {right_data}')

        except Exception as e:
            logger.error(f'Ошибка при поиске в файле {file_name}: {e}')

# TODO--> Нужно доделать передачу названия файла полученного файла для того, чтобы взять в отработку, сейчас функция работает,
# TODO--> но для открытия файла нужен и путь и само название файла
# TODO--> необходимо изменить название модуля так как будут файлы не только PDF
# TODO--> необходимо изменить get_config_value диреткория будет содеражть все файлы




# Пример использования функции
search_terms = {term.lower() for term in {'Протяженность оси трассы', 'СНиП 3.06.04-91', 'ТР ТС 014/2011'}}

def FileChecker():
    extractor = FileChecker(ConfigSettings.get_config_value('xml_zip_local_directory'),
                        ConfigSettings.get_config_value('xml_output_local_directory'),
                        ConfigSettings.get_config_value('pdf_zip_archive_local_directory'),
                        ConfigSettings.get_config_value(
                            'pdf_output_local_directory'))
    extractor.extract_xml()
    extractor.extract_documents()