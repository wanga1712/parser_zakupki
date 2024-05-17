import os
import re
import docx
import pandas as pd
import pdfplumber
import textwrap
import win32com.client as win32
from loguru import logger
import pytesseract
from pdf2image import convert_from_path
import tqdm
from PIL import Image
import cv2
import numpy as np

from config import ConfigSettings

poppler_path = r"C:\Program Files\poppler-24.02.0\Library\bin"
os.environ["PATH"] += os.pathsep + poppler_path

tesseract_path = "C:\Program Files\Tesseract-OCR"
os.environ["PATH"] += os.pathsep + tesseract_path

class DocumentSearcher:
    def __init__(self):
        self.directory = ConfigSettings.get_config_value('unpacked_output_local_directory')

    def format_text(self, text, width=110):
        """Форматирует текст, добавляя переносы строк каждые 'width' символов."""
        return "\n".join(textwrap.wrap(text, width=width))

    def search_documents(self, phrases):
        """Основной метод для обработки директории с файлами."""
        for root, dirs, files in os.walk(self.directory):
            for file in files:
                file_path = os.path.join(root, file)
                file_extension = os.path.splitext(file)[1].lower()
                if file_extension in ('.docx', '.doc'):
                    self.search_in_docx(file_path, phrases)
                elif file_extension in ('.xlsx', '.xls'):
                    self.search_in_excel(file_path, phrases)
                elif file_extension == '.pdf':
                    try:
                        self.search_in_pdf_with_plumber(file_path, phrases)
                    except Exception:
                        self.search_in_pdf_with_ocr(file_path, phrases)

    def read_doc_file(self, file_path):
        """Читает текст из документа Word (.doc) с помощью COM автоматизации."""
        word = win32.gencache.EnsureDispatch('Word.Application')
        word.Visible = False
        doc_text = ""
        try:
            doc = word.Documents.Open(file_path, ReadOnly=True)
            doc_text = doc.Content.Text
            doc.Close(False)
        except Exception as e:
            logger.error(f"Ошибка при чтении файла {file_path}: {e}")
        finally:
            word.Quit()
        return doc_text

    def search_in_docx(self, file_path, phrases):
        """Поиск в документах Word."""
        file_extension = os.path.splitext(file_path)[1].lower()
        try:
            if file_extension == '.docx':
                doc = docx.Document(file_path)
                text_blocks = [para.text for para in doc.paragraphs] + \
                              [' '.join(cell.text for cell in row.cells) for table in doc.tables for row in table.rows]
            elif file_extension == '.doc':
                text_blocks = [self.read_doc_file(file_path)]
            else:
                logger.error(f"Неподдерживаемый формат файла {file_path}")
                return

            for text in text_blocks:
                for phrase in phrases:
                    if phrase.lower() in text.lower():
                        # Найдем индекс начала искомого слова в тексте
                        start_index = text.lower().find(phrase.lower())
                        # Определим начало и конец выделенной области
                        start_slice = max(0, start_index - 30)  # Начало среза, не меньше 0
                        end_slice = min(len(text),
                                        start_index + 30 + len(phrase))  # Конец среза, не больше длины текста
                        # Обрезаем текст до 30 символов перед и после найденного слова
                        sliced_text = text[start_slice:end_slice]
                        logger.info(self.format_text(f'Найдено в {file_path}: "{sliced_text}"'))

        except Exception as e:
            logger.error(f'Ошибка при поиске в файле {file_path}: {e}')

    def search_in_excel(self, file_path, phrases):
        """Поиск в документах Excel."""
        try:
            xls = pd.ExcelFile(file_path)
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name)
                # Очистка от NaN перед созданием строки
                for index, row in df.iterrows():
                    # Заменяем NaN на пустые строки и создаем единую строку
                    row_str = ' '.join(row.fillna('').astype(str))
                    # Удаляем лишние пробелы, которые могли образоваться
                    row_str = re.sub(r'\s+', ' ', row_str).strip()
                    for phrase in phrases:
                        if phrase.lower() in row_str.lower():
                            logger.info(self.format_text(f'Найдено в {file_path} на листе {sheet_name}: "{row_str}"'))
        except Exception as e:
            logger.error(f'Ошибка при поиске в файле {file_path}: {e}')

    def search_in_pdf(self, file_path, phrase):
        """Поиск в PDF документах."""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        # Используем регулярное выражение для поиска контекста вокруг фразы
                        # Слова до и после фразы (\w+\s+){0,10} - до и после до 10 слов
                        pattern = re.compile(r'(\w+\s+){0,10}' + re.escape(phrase) + r'(\s+\w+){0,10}', re.IGNORECASE)
                        results = pattern.findall(text)
                        for result in results:
                            logger.info(
                                f'Найдено в {file_path} на странице {page.page_number}: "{result[0].strip()} {phrase} {result[1].strip()}"')
        except Exception as e:
            raise e  # Передаем исключение дальше

    def search_in_pdf_with_ocr(self, file_path, phrases):
        """Поиск в PDF документах с использованием OCR с отображением прогресса выполнения."""
        logger.info("запустил функцию поиска в отсканированном документе")
        try:
            images = convert_from_path(file_path)
            total_images = len(images)
            found_flag = False  # Флаг для отслеживания наличия найденной фразы
            for i, image in enumerate(tqdm.tqdm(images, desc="OCR прогресс", total=total_images)):
                text = pytesseract.image_to_string(image, lang='rus')  # Указываем язык как 'rus'
                for phrase in phrases:
                    phrase_lower = phrase.lower()
                    text_lower = text.lower()
                    if phrase_lower in text_lower:
                        # Находим индексы начала и конца найденной фразы
                        start_index = text_lower.find(phrase_lower)
                        end_index = start_index + len(phrase_lower)
                        # Определяем индексы начала и конца подстроки для вывода
                        start_substr = max(0, start_index - 40)
                        end_substr = min(len(text), end_index + 60)
                        # Извлекаем подстроку для вывода
                        context = text[start_substr:end_substr]
                        logger.info(f'Найдено в {file_path} на странице {i + 1}: "{context}"')
                        found_flag = True  # Устанавливаем флаг в True при обнаружении фразы
                if found_flag:
                    # Если фраза найдена, выходим из цикла
                    break
        except Exception as e:
            logger.error(f'Ошибка при поиске в файле {file_path} с использованием OCR: {e}')


if __name__ == "__main__":
    searcher = DocumentSearcher()
    phrases = ['светильник', 'светильники', 'светодиод', 'освещение', 'архимет', 'фонарь', 'led', 'лм', 'асуно',
               'система управления', 'автоматическая', 'автоматическая система управления светильниками', 'кулон', 'поламповый контроль', 'прожектор']
    searcher.search_documents(phrases)