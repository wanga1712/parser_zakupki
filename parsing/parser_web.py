import os
import re
import docx
import pandas as pd
import pdfplumber
import textwrap
import win32com.client as win32
from loguru import logger
from config import ConfigSettings


class DocumentSearcher:
    def __init__(self):
        self.directory = ConfigSettings.get_config_value('unpacked_output_local_directory')
        self.checked_files = set()

    def format_text(self, text, width=110):
        """Форматирует текст, добавляя переносы строк каждые 'width' символов."""
        return "\n".join(textwrap.wrap(text, width=width))

    def search_documents(self, phrases):
        """Основной метод для обработки директории с файлами."""
        for root, dirs, files in os.walk(self.directory):
            for file in files:
                file_path = os.path.join(root, file)
                if file_path in self.checked_files:
                    logger.debug(f"Файл {file_path} уже проверен.")
                    continue
                file_extension = os.path.splitext(file)[1].lower()
                if file_extension in ('.docx', '.doc', '.xlsx', '.xls', '.pdf'):
                    found = False
                    if file_extension in ('.docx', '.doc'):
                        found = self.search_in_docx(file_path, phrases)
                    elif file_extension in ('.xlsx', '.xls'):
                        found = self.search_in_excel(file_path, phrases)
                    elif file_extension == '.pdf':
                        found = self.search_in_pdf(file_path, phrases)

                    if found:
                        self.checked_files.add(file_path)

    def read_doc_file(self, file_path):
        """Читает текст из документа Word (.doc) с помощью COM автоматизации."""
        word = win32.gencache.EnsureDispatch('Word.Application')
        word.Visible = False
        doc_text = ""
        try:
            doc = word.Documents.Open(file_path, ReadOnly=True)
            doc_text = doc.Content.Text
            doc.Close(False)
        except Exception as e:
            logger.error(f"Ошибка при чтении файла {file_path}: {e}")
        finally:
            word.Quit()
        return doc_text

    def search_in_docx(self, file_path, phrases):
        """Поиск в документах Word."""
        found = False
        file_extension = os.path.splitext(file_path)[1].lower()
        try:
            if file_extension == '.docx':
                doc = docx.Document(file_path)
                text_blocks = [para.text for para in doc.paragraphs] + \
                              [' '.join(cell.text for cell in row.cells) for table in doc.tables for row in table.rows]
            elif file_extension == '.doc':
                text_blocks = [self.read_doc_file(file_path)]
            else:
                logger.error(f"Неподдерживаемый формат файла {file_path}")
                return found

            for text in text_blocks:
                for phrase in phrases:
                    if phrase.lower() in text.lower():
                        found = True
                        logger.info(self.format_text(f'Найдено в {file_path}: "{text}"'))

        except Exception as e:
            logger.error(f'Ошибка при поиске в файле {file_path}: {e}')
        return found

    def search_in_excel(self, file_path, phrases):
        """Поиск в документах Excel."""
        found = False
        try:
            xls = pd.ExcelFile(file_path)
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name)
                for index, row in df.iterrows():
                    row_str = ' '.join(row.fillna('').astype(str))
                    row_str = re.sub(r'\s+', ' ', row_str).strip()
                    for phrase in phrases:
                        if phrase.lower() in row_str.lower():
                            found = True
                            logger.info(self.format_text(f'Найдено в {file_path} на листе {sheet_name}: "{row_str}"'))
        except Exception as e:
            logger.error(f'Ошибка при поиске в файле {file_path}: {e}')
        return found

    def search_in_pdf(self, file_path, phrase):
        """Поиск в PDF документах."""
        found = False
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pattern = re.compile(r'(\w+\s+){0,10}' + re.escape(phrase) + r'(\s+\w+){0,10}', re.IGNORECASE)
                        results = pattern.findall(text)
                        for result in results:
                            found = True
                            logger.info(
                                f'Найдено в {file_path} на странице {page.page_number}: "{result[0].strip()} {phrase} {result[1].strip()}"')
        except Exception as e:
            logger.error(f'Ошибка при поиске в файле {file_path}: {e}')
        return found


if __name__ == "__main__":
    searcher = DocumentSearcher()
    phrases = ['светильник', 'светильники', 'светильников', 'светодиод', 'освещение', 'светильниками', 'осветить', 'светодиодный']
    searcher.search_documents(phrases)
